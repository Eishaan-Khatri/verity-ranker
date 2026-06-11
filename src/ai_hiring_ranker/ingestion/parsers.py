"""
File parsers for every supported resume and JD format.

Supported formats:
  - .txt / .md  — plain UTF-8 text, read directly
  - .pdf        — extracted via pypdf (text-layer only; no OCR)
  - .docx       — extracted via python-docx

All parsers return a plain str.  Encoding errors are replaced rather than
raised so the pipeline never hard-crashes on a malformed file.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------


def parse_text_file(path: Path) -> str:
    """Read a .txt or .md file, falling back to latin-1 on UTF-8 decode errors."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed for %s; retrying with latin-1", path.name)
        return path.read_text(encoding="latin-1")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def parse_pdf(path: Path) -> str:
    """Extract text from a PDF using pypdf.

    Only the embedded text layer is used — scanned PDFs without an OCR layer
    will return an empty string.  The caller should validate the result.
    """
    try:
        import pypdf  # lazy import so the rest of the pipeline works without it
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF parsing. Install it with: pip install pypdf"
        ) from exc

    pages: list[str] = []
    with pypdf.PdfReader(str(path)) as reader:
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

    full_text = "\n".join(pages)
    if not full_text.strip():
        logger.warning(
            "PDF %s yielded no extractable text. It may be a scanned image PDF.", path.name
        )
    return full_text


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def parse_docx(path: Path) -> str:
    """Extract text from a .docx file using python-docx.

    Paragraph text is joined with newlines.  Table cell text is included after
    all paragraphs so structured resume tables are not silently dropped.
    """
    try:
        import docx  # python-docx exposes the `docx` module
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
        ) from exc

    doc = docx.Document(str(path))

    # Normal paragraph text
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Table cell text (e.g. skills tables, experience grids)
    table_cells: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_cells.append(row_text)

    sections = paragraphs
    if table_cells:
        sections += ["", "--- Table Content ---"] + table_cells

    return "\n".join(sections)


# ---------------------------------------------------------------------------
# Unified dispatcher
# ---------------------------------------------------------------------------


_SUFFIX_MAP: dict[str, callable] = {
    ".txt": parse_text_file,
    ".md": parse_text_file,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
}


def extract_text(path: Path) -> str:
    """Dispatch to the correct parser based on file extension.

    Raises:
        ValueError: if the file extension is not supported.
        FileNotFoundError: if the path does not exist.
    """
    if not path.exists():
        raise FileNotFoundError(f"Input file not found: {path}")

    suffix = path.suffix.lower()
    parser = _SUFFIX_MAP.get(suffix)
    if parser is None:
        raise ValueError(
            f"Unsupported file format '{suffix}' for {path.name}. "
            f"Supported formats: {', '.join(_SUFFIX_MAP)}"
        )

    logger.debug("Parsing %s with %s", path.name, parser.__name__)
    return parser(path)
